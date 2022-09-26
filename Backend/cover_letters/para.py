from docx import Document
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, RGBColor, Pt
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


def add_para(
    content,
    cell=Document(),
    document=Document(),
    align="left",
    font_name="Times New Roman",
    font_size=15,
    font_bold=False,
    font_italic=False,
    font_underline=False,
    color=RGBColor(100, 200, 0),
    keep_together=True,
    keep_with_next=False,
    page_break_before=False,
    widow_control=False,
    line_spacing=1.5,
    columns="1",
    space_before=0,
    space_after=0,
    style_name="Hello",
    image="/home/nevish-302/Downloads/manhwa_suits-2000045.jpg!d.jpg",
):
    # Paragraph 1 text
    paragraph = cell.add_paragraph(content)

    # Configuring paragraph spacing
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_after = Pt(space_after)
    paragraph_format.space_before = Pt(space_before)

    # Configuring line spacing
    paragraph.line_spacing = Inches(line_spacing)

    # Paragraph font configuration
    paragraph.style = document.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = font_name
    font.size = Pt(font_size)
    font.bold = font_bold
    font.italic = font_italic
    font.underline = font_underline
    font.color.rgb = color

    # Paragraph format
    paragraph_format = paragraph.paragraph_format
    paragraph_format.keep_together = keep_together
    paragraph_format.keep_with_next = keep_with_next
    paragraph_format.page_break_before = page_break_before
    paragraph_format.widow_control = widow_control

    # Configuring Alignment

    if align.lower() == "left":
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    elif align.lower() == "center":
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    elif align.lower() == "right":
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    elif align.lower() == "justify":
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    else:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    return paragraph
